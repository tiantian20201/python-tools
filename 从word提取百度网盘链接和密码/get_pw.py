"""
# Document 还有添加标题、分页、段落、图片、章节等方法，说明如下
  |  add_heading(self, text='', level=1)
  |      Return a heading paragraph newly added to the end of the document,
  |      containing *text* and having its paragraph style determined by
  |      *level*. If *level* is 0, the style is set to `Title`. If *level* is
  |      1 (or omitted), `Heading 1` is used. Otherwise the style is set to
  |      `Heading {level}`. Raises |ValueError| if *level* is outside the
  |      range 0-9.
  |
  |  add_page_break(self)
  |      Return a paragraph newly added to the end of the document and
  |      containing only a page break.
  |
  |  add_paragraph(self, text='', style=None)
  |      Return a paragraph newly added to the end of the document, populated
  |      with *text* and having paragraph style *style*. *text* can contain
  |      tab (``\t``) characters, which are converted to the appropriate XML
  |      form for a tab. *text* can also include newline (``\n``) or carriage
  |      return (``\r``) characters, each of which is converted to a line
  |      break.
  |
  |  add_picture(self, image_path_or_stream, width=None, height=None)
  |      Return a new picture shape added in its own paragraph at the end of
  |      the document. The picture contains the image at
  |      *image_path_or_stream*, scaled based on *width* and *height*. If
  |      neither width nor height is specified, the picture appears at its
  |      native size. If only one is specified, it is used to compute
  |      a scaling factor that is then applied to the unspecified dimension,
  |      preserving the aspect ratio of the image. The native size of the
  |      picture is calculated using the dots-per-inch (dpi) value specified
  |      in the image file, defaulting to 72 dpi if no value is specified, as
  |      is often the case.
  |
  |  add_section(self, start_type=2)
  |      Return a |Section| object representing a new section added at the end
  |      of the document. The optional *start_type* argument must be a member
  |      of the :ref:`WdSectionStart` enumeration, and defaults to
  |      ``WD_SECTION.NEW_PAGE`` if not provided.
  |
  |  add_table(self, rows, cols, style=None)
  |      Add a table having row and column counts of *rows* and *cols*
  |      respectively and table style of *style*. *style* may be a paragraph
  |      style object or a paragraph style name. If *style* is |None|, the
  |      table inherits the default table style of the document.
  |
  |  save(self, path_or_stream)
  |      Save this document to *path_or_stream*, which can be eit a path to
  |      a filesystem location (a string) or a file-like object.

"""

# 读取docx中的文本代码示例
import docx
import re
#https://python-docx.readthedocs.io/en/latest/user/quickstart.html#opening-a-document

# 获取文档table对象

with open("pws","w",encoding="utf-8") as f2:
	file = docx.Document("2000g.docx")
	for table in file.tables:
		for row in table.rows:
			for cell in row.cells:
				if len(re.findall(r"^[a-z0-9A-Z]{4}$",cell.text)) > 0: # 过滤密码
					# print(cell.text)
					f2.write(cell.text+"\n")

# print("段落数:" + str(len(file.paragraphs)))  # 段落数，每个回车隔离一段

# 输出段落编号及段落内容
# for i in range(len(file.paragraphs)):
#     print("第" + str(i) + "段的内容是：" + file.paragraphs[i].text)
